import io
import docx
import pytest
from unittest.mock import AsyncMock, MagicMock
from app.services.ingestor_service import IngestorService

@pytest.mark.asyncio
async def test_extraction_logic():
    # 1. Create a dummy DOCX in memory
    doc = docx.Document()
    doc.add_paragraph("Hello World")
    doc.add_paragraph("This is a test document.")
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_bytes = file_stream.getvalue()
    
    # 2. Setup Service with mocks
    repo = MagicMock()
    s3 = MagicMock()
    sqs = MagicMock()
    
    service = IngestorService(repo, s3, sqs)
    
    # 3. Test extraction
    text = service.extract_text_from_docx(file_bytes)
    
    assert "Hello World" in text
    assert "This is a test document." in text
    assert "\n" in text

@pytest.mark.asyncio
async def test_process_batch_success():
    # Setup
    repo = AsyncMock()
    s3 = AsyncMock()
    sqs = AsyncMock()
    
    # Create dummy docx bytes
    doc = docx.Document()
    doc.add_paragraph("Test content")
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_bytes = file_stream.getvalue()
    
    # Mock return values
    record = MagicMock()
    record.doc_id = "123"
    record.user_id = "user1"
    record.file_name = "test.docx"
    
    repo.claim_pending_documents.return_value = [record]
    s3.download_file.return_value = file_bytes
    
    service = IngestorService(repo, s3, sqs)
    
    # Execute
    processed = await service.process_batch(limit=1)
    
    # Verify
    assert processed == 1
    repo.claim_pending_documents.assert_called_once()
    s3.download_file.assert_called_once_with("123_test.docx")
    sqs.send_task.assert_called_once()
    repo.update_status.assert_called_with("123", "Ingested")
