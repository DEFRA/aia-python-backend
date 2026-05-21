import io

import docx
import pytest
from app.services.ingestor_service import IngestorService


@pytest.mark.asyncio
async def test_extraction_logic():
    doc = docx.Document()
    doc.add_paragraph("Hello World")
    doc.add_paragraph("This is a test document.")

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_bytes = file_stream.getvalue()

    service = IngestorService()

    text = service.extract_text_from_docx(file_bytes)

    assert "Hello World" in text
    assert "This is a test document." in text
    assert "\n" in text


def test_extract_raises_on_empty_bytes():
    service = IngestorService()
    with pytest.raises(ValueError, match="empty"):
        service.extract_text_from_docx(b"")


def test_extract_raises_on_invalid_bytes():
    service = IngestorService()
    with pytest.raises(ValueError, match="Failed to open"):
        service.extract_text_from_docx(b"not-a-docx")
