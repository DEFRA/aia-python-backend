import io

import docx

from app.utils.logger import get_logger

logger = get_logger(__name__)


class IngestorService:
    def extract_text_from_docx(self, file_bytes: bytes) -> str:
        """
        Extracts all readable text from a DOCX file, including paragraphs and tables.

        Raises ValueError for empty input, unreadable files, or documents with no text.
        """
        if not file_bytes:
            raise ValueError("Cannot extract text: file content is empty.")

        try:
            doc = docx.Document(io.BytesIO(file_bytes))
        except Exception as exc:
            raise ValueError(f"Failed to open DOCX document: {exc}") from exc

        lines: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        lines.append(text)

        if not lines:
            raise ValueError("DOCX file contains no extractable text content.")

        extracted = "\n".join(lines)
        logger.debug(
            "Extracted %d chars from DOCX (%d bytes input)",
            len(extracted),
            len(file_bytes),
        )
        return extracted
